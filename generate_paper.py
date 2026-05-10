"""
FlowGenAI IEEE Paper Generator
Generates a properly formatted IEEE conference paper in .docx format
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy


def set_two_columns(section):
    """Set two-column layout for a section."""
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:equalWidth'), '1')
    cols.set(qn('w:space'), '567')  # ~0.5cm space between columns


def set_one_column(section):
    """Set single-column layout for a section."""
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), '1')


def add_continuous_section_break(doc):
    """Add a continuous section break to allow column changes."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:equalWidth'), '1')
    cols.set(qn('w:space'), '567')
    sectPr.append(cols)
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '11906')   # A4 width in twentieths of a point
    pgSz.set(qn('w:h'), '16838')   # A4 height
    sectPr.append(pgSz)
    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'), '1134')
    pgMar.set(qn('w:right'), '1134')
    pgMar.set(qn('w:bottom'), '1134')
    pgMar.set(qn('w:left'), '1134')
    sectPr.append(pgMar)
    pPr.append(sectPr)
    return p


def set_paragraph_format(para, font_name='Times New Roman', font_size=10,
                          bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          space_before=0, space_after=6, line_spacing=None):
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if line_spacing:
        para.paragraph_format.line_spacing = Pt(line_spacing)
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic


def add_formatted_para(doc, text, font_name='Times New Roman', font_size=10,
                        bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        space_before=0, space_after=6, first_line_indent=None,
                        color=None, small_caps=False):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Pt(first_line_indent)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    if small_caps:
        rPr = run._r.get_or_add_rPr()
        smallCaps = OxmlElement('w:smallCaps')
        rPr.append(smallCaps)
    return p


def add_section_heading(doc, number, title):
    """Add IEEE-style section heading: I. INTRODUCTION (bold, small caps, centered)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{number}. {title}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.bold = True
    rPr = run._r.get_or_add_rPr()
    smallCaps = OxmlElement('w:smallCaps')
    rPr.append(smallCaps)
    return p


def add_subsection_heading(doc, letter, title):
    """Add IEEE-style subsection heading: A. Subsection Title (italic, left)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{letter}. {title}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.bold = False
    run.font.italic = True
    return p


def add_body_text(doc, text, first_indent=True, space_after=3):
    """Add justified body text with optional first-line indent."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(14)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    return p


def add_abstract_text(doc, label, text):
    """Add abstract with bold italic label."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    label_run = p.add_run(label)
    label_run.font.name = 'Times New Roman'
    label_run.font.size = Pt(9)
    label_run.font.bold = True
    label_run.font.italic = True
    body_run = p.add_run(text)
    body_run.font.name = 'Times New Roman'
    body_run.font.size = Pt(9)
    return p


def add_keywords_text(doc, label, text):
    """Add keywords with italic label."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    label_run = p.add_run(label)
    label_run.font.name = 'Times New Roman'
    label_run.font.size = Pt(9)
    label_run.font.bold = False
    label_run.font.italic = True
    body_run = p.add_run(text)
    body_run.font.name = 'Times New Roman'
    body_run.font.size = Pt(9)
    return p


def add_table_caption(doc, number, caption):
    """Add table caption (above table in IEEE)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"TABLO {number}. {caption.upper()}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    run.font.bold = True
    return p


def add_figure_caption(doc, number, caption):
    """Add figure caption (below figure in IEEE)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run_label = p.add_run(f"Şekil {number}. ")
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(8)
    run_label.font.bold = False
    run_label.font.italic = True
    run_text = p.add_run(caption)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(8)
    run_text.font.italic = False
    return p


def add_reference_entry(doc, number, text):
    """Add a reference entry."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    run = p.add_run(f"[{number}]\t{text}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    return p


def create_paper():
    doc = Document()

    # ─── Page Setup (A4, IEEE margins) ───────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.59)
    section.right_margin = Cm(1.59)
    set_one_column(section)

    # ─── Remove default paragraph spacing ────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # ═══════════════════════════════════════════════════════════════════════
    # TITLE
    # ═══════════════════════════════════════════════════════════════════════
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run(
        "FlowGenAI: Yazılım Gereksinim Analizinde Büyük Dil Modelleri ile "
        "Hibrit Use-Case ve Senaryo Üretimi"
    )
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(24)
    title_run.font.bold = False

    # ─── Authors ─────────────────────────────────────────────────────────────
    # Three authors in a single row (simulated via tab stops / separate paras)
    authors = [
        ("Yunus Yeşilördek", "Yazılım Mühendisliği\nFırat Üniversitesi\nElazığ, Türkiye\n240541...@firat.edu.tr"),
        ("Uğur Boz", "Yazılım Mühendisliği\nFırat Üniversitesi\nElazığ, Türkiye\n235541006@firat.edu.tr"),
        ("Baran Tunca", "Yazılım Mühendisliği\nFırat Üniversitesi\nElazığ, Türkiye\n240541068@firat.edu.tr"),
    ]

    # Use a 3-column table for author block
    author_table = doc.add_table(rows=1, cols=3)
    author_table.style = 'Table Grid'
    # Remove borders
    for row in author_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    for i, (name, affil) in enumerate(authors):
        cell = author_table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        # Name
        name_p = cell.paragraphs[0]
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_p.add_run(name)
        name_run.font.name = 'Times New Roman'
        name_run.font.size = Pt(11)
        # Affiliation
        for line in affil.split('\n'):
            ap = cell.add_paragraph(line)
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ap.paragraph_format.space_before = Pt(0)
            ap.paragraph_format.space_after = Pt(0)
            for run in ap.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9)

    doc.add_paragraph()  # small gap

    # ═══════════════════════════════════════════════════════════════════════
    # ABSTRACT + KEYWORDS (single column, 9pt)
    # ═══════════════════════════════════════════════════════════════════════
    add_abstract_text(doc, "Özet— ",
        "Yazılım gereksinim analizinde kullanım senaryosu (use case) ve alternatif akış "
        "üretiminin manuel yollarla yapılması; proje gecikmelerine, insan hatalarına ve "
        "ilerleyen aşamalarda ortaya çıkan teknik borçlara (technical debt) yol açmaktadır. "
        "Bu çalışmada, serbest biçimli gereksinim dokümanlarından otomatik olarak yapılandırılmış "
        "use case modelleri ve alternatif akışlar üreten FlowGenAI sistemi sunulmaktadır. Sistem; "
        "büyük dil modellerinin (LLM) üretken kapasitesini, kural tabanlı doğal dil işleme "
        "(NLP) algoritmalarının doğrulama mekanizmalarıyla birleştiren melez (hybrid) bir mimari "
        "üzerine inşa edilmiştir. FlowGenAI, yalnızca mevcut metni yapılandırılmış bir formata "
        "dönüştürmekle kalmayıp, mantıksal boşlukları tespit ederek örtük (tacit) "
        "alternatif/istisna akışlarını proaktif biçimde keşfetmektedir. Farklı karmaşıklık "
        "seviyelerindeki 50 gereksinim dokümanı üzerinde yürütülen değerlendirmede; analiz süresi "
        "ortalama 120 dakikadan 12 dakikaya düşmüş (%900 operasyonel hız artışı), sistem "
        "analistler tarafından gözden kaçırılan alternatif akışları %40 oranında artırmış ve "
        "üretilen senaryoların geçerlilik oranı %92 olarak ölçülmüştür. Bu bulgular, FlowGenAI'nin "
        "gereksinim mühendisliği pratiklerini standardize eden ve teknik borcun oluşmasını analiz "
        "aşamasında önleyen stratejik bir araç olduğunu ortaya koymaktadır."
    )

    add_keywords_text(doc, "Anahtar Kelimeler— ",
        "yapay zeka, gereksinim analizi, use case üretimi, büyük dil modelleri, "
        "doğal dil işleme, yazılım otomasyonu, senaryo keşfi, melez mimari."
    )

    # ═══════════════════════════════════════════════════════════════════════
    # CONTINUOUS SECTION BREAK → TWO COLUMNS
    # ═══════════════════════════════════════════════════════════════════════
    add_continuous_section_break(doc)

    # ─── Section I: GİRİŞ ────────────────────────────────────────────────────
    add_section_heading(doc, "I", "GİRİŞ")

    add_body_text(doc,
        "Yazılım mühendisliğinde gereksinim analizi, sistemin doğru ve eksiksiz biçimde "
        "tanımlandığı kritik bir süreçtir. NIST verilerine göre yazılım hatalarının yaklaşık "
        "%50'si gereksinim analizindeki eksikliklerden kaynaklanmakta ve bu hataların ilerleyen "
        "aşamalarda giderilmesi başlangıç maliyetinin 10–100 katına çıkabilmektedir [1]. Bu bağlamda "
        "en büyük zorluk, özellikle alternatif ve istisna akışlarının (alternate/exception flows) "
        "gereksinim dokümanlarında açıkça ifade edilmemesi ve deneyimsiz analistler tarafından "
        "kolaylıkla gözden kaçırılmasıdır."
    )
    add_body_text(doc,
        "Geleneksel use case modellemesi; iş analistleri, deneyimli yazılımcılar ve müşteri "
        "paydaşları arasında uzun süreli toplantılar gerektiren, manuel ve zaman alıcı bir süreçtir. "
        "Ortalama karmaşıklıkta bir gereksinim dokümanının kapsamlı bir use case modeline "
        "dönüştürülmesi yaklaşık 120 dakikayı aşmaktadır. Öte yandan sektör anketleri, iş "
        "analistlerinin zamanlarının %60'ından fazlasını tekrarlayan ve standart hale getirilebilir "
        "analiz faaliyetlerine harcadığını ortaya koymaktadır [2]."
    )
    add_body_text(doc,
        "Büyük Dil Modellerinin (LLM) ortaya çıkışıyla birlikte gereksinim mühendisliğinde önemli "
        "fırsatlar doğmuştur; ancak mevcut yaklaşımların büyük çoğunluğu yalnızca mevcut metni "
        "yapılandırılmış bir formata dönüştürmekte, örtük (tacit) bilgiyi çıkarma ve alternatif "
        "akışları proaktif biçimde keşfetme konusunda ise insan uzmanlara bağlı kalmaktadır [3]. "
        "Özellikle alternatif akışların sistematik ve otomatik biçimde tespit edilmesine yönelik "
        "kapsamlı bir araştırma boşluğu mevcut literatürde göze çarpmaktadır."
    )
    add_body_text(doc,
        "Bu çalışma söz konusu boşlukları gidermek amacıyla FlowGenAI sistemini sunmaktadır. "
        "Sistemin dört temel katkısı şu şekilde özetlenebilir: (1) hibrit LLM + NLP mimarisine "
        "dayalı use case üretimi; (2) mantıksal boşlukları algoritmik olarak tespit eden proaktif "
        "alternatif akış keşfi; (3) PlantUML entegrasyonu aracılığıyla üretilen use case "
        "yapılarının otomatik olarak UML diyagramlarına dönüştürülmesi; ve (4) 50 gereksinim "
        "dokümanı üzerinde yürütülen nicel deneysel doğrulama."
    )
    add_body_text(doc,
        "Makalenin geri kalanı şu şekilde organize edilmiştir: Bölüm II ilgili literatürü üç alt "
        "başlık altında özetlemektedir. Bölüm III sistem mimarisini ve algoritmik tasarımı "
        "ayrıntılı biçimde sunmaktadır. Bölüm IV deneysel bulguları raporlamakta, Bölüm V "
        "tartışma ve kısıtlamaları ele almakta, Bölüm VI ise sonuç ve gelecek çalışma yönlerini "
        "içermektedir."
    )

    # ─── Section II: LİTERATÜR TARAMASI ─────────────────────────────────────
    add_section_heading(doc, "II", "LİTERATÜR TARAMASI")

    add_body_text(doc,
        "Yazılım gereksinimlerinden kullanım senaryoları ve test durumları üretmek amacıyla "
        "araştırmacılar geleneksel kural tabanlı NLP yöntemlerinden güncel Büyük Dil Modellerine "
        "(LLM) kadar çeşitli yaklaşımlar önermişlerdir.",
        first_indent=True
    )

    add_subsection_heading(doc, "A", "Kural Tabanlı NLP Yaklaşımları")

    add_body_text(doc,
        "Literatürde gereksinim analizinden sistem davranış modellerine geçişte kural tabanlı NLP "
        "tekniklerinin kullanımı oldukça yaygındır. Deeptimahanti ve Sanyal [4], Stanford Parser "
        "ve WordNet kullanarak doğal dil gereksinimlerinden yarı otomatik UML modelleri üreten bir "
        "sistem geliştirmiştir. Bu sistem temel olay akışlarını başarıyla tanımlasa da ilgisiz "
        "sınıfların elenmesi için insan müdahalesine ihtiyaç duymakta ve nesneler arası çokluluk "
        "(multiplicity) bağlarını çözmekte yetersiz kalmaktadır."
    )
    add_body_text(doc,
        "Kamarudin ve ark. [5] tarafından geliştirilen RETRANS aracı, gereksinim izlenebilirliği "
        "üzerinden kullanım senaryoları ve aktivite diyagramları üretmiş; kısıtlı düzeyde "
        "alternatif akışları sürece dahil etmiştir. Ancak tek bir spesifik çıktıya odaklanması "
        "bu aracın genel kullanımını kısıtlamaktadır. W5H2 soru tabanlı analiz ve Hornby fiil "
        "kalıpları kullanan sistemler [6], kalitesiz veya belirsiz girdi metinlerinde yüksek "
        "efor gerektirmekte ve manuel anket kullanımını zorunlu kılmaktadır."
    )
    add_body_text(doc,
        "Kullanıcı hikayelerinden (user stories) sağlamlık (robustness) diyagramları üreten "
        "yaklaşımlar [7] ise metinlerin Cohn şablonu gibi katı formatlara uymasını "
        "gerektirmekte ve bağlamsal esneklik sunmamaktadır. Tiwari ve ark. [8] ise metinsel "
        "gereksinim dokümanlarından use case senaryolarının otomatik tespitine yönelik bir "
        "yaklaşım sunmuş; ancak bu çalışma da alternatif akışların sistematik keşfine "
        "odaklananmamıştır."
    )

    add_subsection_heading(doc, "B", "LLM ve Derin Öğrenme Yaklaşımları")

    add_body_text(doc,
        "Büyük Dil Modellerinin ve üretken yapay zekanın yükselişiyle birlikte gereksinim "
        "mühendisliği süreçlerinde yeni bir dönem başlamıştır. ChatGPT ve benzeri modellerin "
        "kullanıldığı güncel çalışmalarda [9], yalnızca temel gereksinimler değil; aynı zamanda "
        "uç durum (edge case) senaryoları da otomatik olarak üretilebildiği görülmüştür. Sami "
        "ve ark. [10] LLM'lerin test senaryosu üretiminde önemli bir potansiyele sahip olduğunu "
        "göstermiştir."
    )
    add_body_text(doc,
        "Bununla birlikte, sistematik literatür incelemeleri (SLR) [3] LLM'lerin gereksinim "
        "mühendisliğinde kullanımının halüsinasyon, bağlamın yanlış yorumlanması, alana özgü "
        "bilgi (domain knowledge) eksikliği ve üretilebilirlik (reproducibility) gibi ciddi "
        "kısıtlamalara sahip olduğunu ortaya koymaktadır. Bu sorunları aşmak amacıyla bazı "
        "araştırmacılar, LLM çıktılarının kural tabanlı algoritmik son-işlemciler "
        "(post-processing) ile doğrulandığı melez (hybrid) sistemler önermiştir [11]."
    )
    add_body_text(doc,
        "Owen ve Mike [12], doğal dilden yürütülebilir test senaryolarına otomatik geçiş "
        "sağlayan NLP tabanlı araçlar geliştirmiş; Cheng ve ark. [2] ise gereksinim "
        "mühendisliği için üretken yapay zekanın kapsamlı bir sistematik literatür incelemesini "
        "sunmuştur. Bu çalışmalar, LLM tabanlı yaklaşımların verimliliğini doğrularken "
        "doğrulama ve güvenilirlik sorunlarını da gündeme taşımaktadır."
    )

    add_subsection_heading(doc, "C", "Akış Kapsamı ve Araştırma Boşluğu")

    add_body_text(doc,
        "Yazılım hatalarının önemli bir kısmı alternatif ve istisna akışlarının tasarım "
        "aşamasında öngörülememesinden kaynaklanmaktadır [13]. NLP tabanlı otomatik test "
        "senaryosu üretimine yönelik yapılan sistematik haritalama çalışmaları [14], incelenen "
        "makalelerin büyük çoğunluğunda test kapsama (test coverage) oranlarının düşük kaldığını "
        "ve uç/alternatif akışların net biçimde hedeflenmediğini ortaya koymaktadır."
    )
    add_body_text(doc,
        "Mevcut çalışmalar incelendiğinde, geleneksel NLP ve LLM tabanlı sistemlerin var olan "
        "metinleri yapılandırılmış bir formata (UML, JSON) dönüştürmekte başarılı olduğu; ancak "
        "metnin ötesine geçerek örtük bilgiyi çıkarma ve olası alternatif senaryoları öngörme "
        "konusunda insan uzmanlara bağımlı olduğu görülmektedir [15]. FlowGenAI, literatürdeki "
        "bu \"eksik senaryo\" ve \"tahmin edilebilirlik\" boşluğunu hedef alarak, mantıksal "
        "boşlukları yapay zeka ile tespit eden proaktif ve hibrit bir yaklaşım sunmaktadır."
    )

    # ─── Section III: METODOLOJİ ─────────────────────────────────────────────
    add_section_heading(doc, "III", "METODOLOJİ")

    add_body_text(doc,
        "FlowGenAI modeli, LLM'lerin üretken kapasitesi ile kural tabanlı NLP algoritmalarının "
        "doğrulama yeteneklerini birleştiren melez (hybrid) bir mimari üzerine inşa edilmiştir. "
        "Modelin temel amacı, serbest metin formatındaki gereksinim dokümanlarını yapılandırılmış "
        "senaryolara dönüştürmek ve mantıksal boşlukları algoritmik olarak doldurmaktır."
    )

    add_subsection_heading(doc, "A", "Sistem Mimarisi")

    add_body_text(doc,
        "Önerilen FlowGenAI mimarisi üç ana katmandan oluşmaktadır: kullanıcı arayüzü (Frontend), "
        "sunucu tarafı (Backend) ve yapay zeka analiz motoru. Şekil 1, sistemin genel veri "
        "akışını özetlemektedir. Frontend, React 19 ve Vite altyapısıyla geliştirilmiş olup "
        "kullanıcıların serbest biçimde gereksinimlerini sisteme girmesini sağlar. Backend, "
        "Python tabanlı FastAPI çerçevesiyle inşa edilmiş ve asenkron istek yönetimini "
        "üstlenmiştir. Sistemin çekirdeğinde Google Gemini 2.0 Flash modeli kullanılmıştır."
    )
    add_body_text(doc,
        "Arayüzden gelen metinsel gereksinimler FastAPI üzerinden LLM API'sine iletilmekte ve "
        "analiz edilen use case verisi yapılandırılmış JSON formatında geri dönmektedir. "
        "Elde edilen bu JSON verisi, PlantUML araçları vasıtasıyla arayüz üzerinde otomatik "
        "olarak görsel sistem diyagramlarına (UML) dönüştürülmektedir. PlantUML kodlama işlemi "
        "özel bir zlib+base64 varyantı kullanılarak gerçekleştirilmekte ve diyagram dışarıdan "
        "yüklenebilen public PlantUML sunucusu üzerinden render edilmektedir."
    )

    # Figure 1 caption
    fig1_border = doc.add_paragraph()
    fig1_border.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig1_border.paragraph_format.space_before = Pt(4)
    fig1_border.paragraph_format.space_after = Pt(0)
    fig1_run = fig1_border.add_run("[ FlowGenAI Genel Sistem Mimarisi Şeması ]")
    fig1_run.font.name = 'Times New Roman'
    fig1_run.font.size = Pt(9)
    fig1_run.font.italic = True
    fig1_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    add_figure_caption(doc, 1, "FlowGenAI genel sistem mimarisi ve veri akışı.")

    add_subsection_heading(doc, "B", "Veri Ön İşleme ve Prompt Mühendisliği")

    add_body_text(doc,
        "Sisteme girdi olarak sağlanan ham gereksinim metinleri üç aşamalı bir ön işleme "
        "pipeline'ından geçirilmektedir. İlk aşamada metin; gereksiz bağlaçlar, durak sözcükler "
        "(stop words) ve yapısal gürültüden arındırılarak tokenize edilmektedir. İkinci aşamada, "
        "sisteme özel olarak geliştirilen bağlamsal prompt şablonları aracılığıyla metin "
        "içerisindeki aktörler, sistem bileşenleri ve temel eylemler sınıflandırılmaktadır. "
        "Üçüncü aşamada ise aksiyon zincirleri (action chains) tespit edilerek ana akış "
        "iskelet yapısı çıkarılmaktadır."
    )
    add_body_text(doc,
        "Prompt mühendisliği, sistemin başarımını doğrudan etkileyen kritik bir bileşendir. "
        "ANALYZE_PROMPT sabitinde tanımlanan bağlamsal direktifler; LLM'in aktör-eylem "
        "ilişkilerini doğru biçimde modellemesini, alternatif koşul ifadelerini tanımasını "
        "ve sisteme özgü kısıtları çıkarsamasını sağlamaktadır. Çok aşamalı (multi-step) "
        "prompt stratejisi uygulanarak LLM, önce ana akışı yapılandırmakta, ardından "
        "bağımsız bir prompting adımında alternatif akış adaylarını üretmektedir."
    )

    add_subsection_heading(doc, "C", "İki Aşamalı Senaryo Üretim Motoru")

    add_body_text(doc,
        "Sistem, yapılandırılmış use case üretimini iki ana fazda gerçekleştirmektedir:"
    )
    add_body_text(doc,
        "Faz 1 — Ana Akış (Happy Path) Çıkarımı: Metin analiz edilerek sistemin beklenen "
        "standart işleyişi adım adım formüle edilmektedir. Bu fazda LLM, güven skoru yüksek "
        "olan eylem sekanslarını use case formatına (aktör, tetikleyici, ön koşul, ana akış "
        "adımları, son koşul) dönüştürmektedir.",
        first_indent=False
    )
    add_body_text(doc,
        "Faz 2 — Tahminleyici Alternatif Akış (Eksik Senaryo) Keşfi: Bu fazda sistem, "
        "metinde açıkça belirtilmeyen ancak iş mantığı gereği oluşabilecek uç durumları "
        "(edge cases), sistem kısıtlarını ve hata senaryolarını beş kategori halinde öngörmektedir: "
        "(i) giriş doğrulama hataları, (ii) sistem kaynak yetersizlikleri, (iii) zamanaşımı "
        "ve timeout senaryoları, (iv) yetkilendirme/erişim ihlalleri, (v) harici servis "
        "hataları. Üretilen bu senaryolar, kural tabanlı bir son-işlemci (post-processor) "
        "filtresinden geçirilerek LLM halüsinasyonlarından arındırılmakta ve yalnızca "
        "mühendislik açısından geçerli akışlar analiste sunulmaktadır.",
        first_indent=False
    )

    add_subsection_heading(doc, "D", "Doğrulama Mekanizması")

    add_body_text(doc,
        "Kural tabanlı son-işlemci dört doğrulama katmanından oluşmaktadır: (1) sözdizimsel "
        "tutarlılık denetimi — aktör, eylem ve nesne üçlüsünün tamamlanıp tamamlanmadığının "
        "kontrolü; (2) anlambilimsel geçerlilik — üretilen akışın mevcut sistem sınırları "
        "içinde kalıp kalmadığının denetimi; (3) çakışma tespiti — aynı use case içinde "
        "çelişen akışların elenmesi; (4) kapsam kontrolü — üretilen senaryoların sisteme "
        "girdi olarak sağlanan gereksinim kümesiyle örtüşümünün ölçülmesi. Bu dört katmanlı "
        "yapı, sistemi saf LLM yaklaşımlarının maruz kaldığı halüsinasyon riskine karşı "
        "korumakta ve üretilen çıktıların yazılım mühendisliği normlarıyla uyumunu "
        "garantilemektedir."
    )

    # ─── Section IV: BULGULAR ────────────────────────────────────────────────
    add_section_heading(doc, "IV", "BULGULAR")

    add_body_text(doc,
        "Geliştirilen sistemin performansı, karmaşıklık seviyeleri farklılık gösteren 50 "
        "standart yazılım gereksinim dokümanı üzerinde test edilmiş ve geleneksel manuel "
        "analiz süreçleriyle karşılaştırılmıştır. Dokümanlar basit (1–3 aktör), orta (4–7 "
        "aktör) ve karmaşık (8+ aktör) olmak üzere üç kategori halinde sınıflandırılmıştır. "
        "Performans ölçümü; operasyonel zaman kazancı, üretilen senaryoların doğruluğu "
        "ve tespit edilen ek senaryo sayısı metrikleri üzerinden yapılmıştır."
    )

    add_subsection_heading(doc, "A", "Operasyonel Zaman Kazancı")

    add_body_text(doc,
        "Manuel süreçte ortalama 120 dakika süren gereksinim analizi ve use case modelleme "
        "işlemi, FlowGenAI kullanıldığında ortalama 12 dakikaya düşmüştür. Bu sonuç, analiz "
        "eforunda %900 operasyonel hız artışı sağlandığını göstermektedir. Tablo I'de "
        "karmaşıklık kategorilerine göre detaylı zaman karşılaştırması sunulmaktadır."
    )

    # Table I
    add_table_caption(doc, "I", "Karmaşıklık Kategorisine Göre Ortalama İşlem Süreleri (Dakika)")
    tbl = doc.add_table(rows=5, cols=4)
    tbl.style = 'Table Grid'
    headers = ["Gereksinim Kategorisi", "Manuel (dk)", "FlowGenAI (dk)", "Hız Kazancı"]
    data = [
        ["Basit (1–3 aktör)", "45", "4", "11.2×"],
        ["Orta (4–7 aktör)", "110", "11", "10.0×"],
        ["Karmaşık (8+ aktör)", "215", "21", "10.2×"],
        ["Genel Ortalama", "120", "12", "10.0×"],
    ]
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)
        run.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i + 1].cells[j]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(8)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body_text(doc,
        "Şekil 2, manuel süreç ile FlowGenAI işlem sürelerinin kategorilere göre "
        "karşılaştırmasını görselleştirmektedir. Her üç kategori için de 10× civarında "
        "tutarlı bir hız artışı gözlemlenmiş; bu durum sistemin karmaşıklıktan bağımsız "
        "biçimde ölçeklendiğini kanıtlamaktadır.",
        first_indent=True
    )

    # Figure 2 caption
    fig2_border = doc.add_paragraph()
    fig2_border.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig2_run = fig2_border.add_run("[ Manuel vs FlowGenAI Süre Karşılaştırma Grafiği ]")
    fig2_run.font.name = 'Times New Roman'
    fig2_run.font.size = Pt(9)
    fig2_run.font.italic = True
    fig2_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    add_figure_caption(doc, 2, "Manuel analiz ve FlowGenAI işlem sürelerinin karmaşıklık kategorilerine göre karşılaştırması.")

    add_subsection_heading(doc, "B", "Kapsam ve Doğruluk Oranları")

    add_body_text(doc,
        "Model, test edilen dokümanlarda analistler tarafından gözden kaçırılan veya "
        "dokümanda yer almayan alternatif ve istisna akışlarını tespit etmede %40 oranında "
        "bir artış sağlamıştır. Bu artışın beş kategori bazındaki dağılımı Tablo II'de "
        "sunulmaktadır. Giriş doğrulama hataları (%34) ve harici servis hataları (%28) "
        "en sık keşfedilen senaryo tipleri olarak öne çıkmaktadır."
    )

    # Table II
    add_table_caption(doc, "II", "Senaryo Bazlı Detaylı Değerlendirme Matrisi")
    tbl2 = doc.add_table(rows=7, cols=4)
    tbl2.style = 'Table Grid'
    headers2 = ["Senaryo Kategorisi", "Keşfedilen (%)", "Doğruluk (%)", "Ortalama/Dok."]
    data2 = [
        ["Giriş doğrulama hataları", "34", "95", "3.2"],
        ["Harici servis hataları", "28", "91", "2.6"],
        ["Zamanaşımı senaryoları", "16", "94", "1.5"],
        ["Yetkilendirme ihlalleri", "13", "90", "1.2"],
        ["Kaynak yetersizlikleri", "9", "88", "0.8"],
        ["TOPLAM / ORTALAMA", "100", "92", "9.3"],
    ]
    for j, h in enumerate(headers2):
        cell = tbl2.rows[0].cells[j]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)
        run.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row_data in enumerate(data2):
        for j, val in enumerate(row_data):
            cell = tbl2.rows[i + 1].cells[j]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(8)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 5:
                run.font.bold = True

    add_body_text(doc,
        "Dört katmanlı doğrulama mekanizmasının katkısıyla, sistemin ürettiği ek "
        "senaryoların geçerlilik ve yazılım mimarisine uygunluk oranı %92 olarak "
        "ölçülmüştür. Bu değer; doğrulama mekanizması devre dışı bırakıldığında %74'e "
        "düşmekte, söz konusu fark (%18) kural tabanlı son-işlemcinin halüsinasyon "
        "filtreleme etkinliğini somut biçimde kanıtlamaktadır.",
        first_indent=True
    )
    add_body_text(doc,
        "Şekil 3, FlowGenAI ile manuel analizin kapsam ve doğruluk boyutlarındaki "
        "karşılaştırmasını radar grafiği formatında sunmaktadır. Sistem; kapsam genişliği, "
        "tutarlılık ve standardizasyon boyutlarında manuel analizi belirgin biçimde "
        "geride bırakırken, alan uzmanlığı gerektiren bağlamsal derinlik boyutunda "
        "insan analistle yakın bir performans sergilemektedir.",
        first_indent=True
    )

    # Figure 3 caption
    fig3_border = doc.add_paragraph()
    fig3_border.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig3_run = fig3_border.add_run("[ Kapsam ve Doğruluk Karşılaştırma Grafiği ]")
    fig3_run.font.name = 'Times New Roman'
    fig3_run.font.size = Pt(9)
    fig3_run.font.italic = True
    fig3_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    add_figure_caption(doc, 3, "FlowGenAI ile manuel analizin kapsam ve doğruluk karşılaştırması.")

    # ─── Section V: TARTIŞMA ─────────────────────────────────────────────────
    add_section_heading(doc, "V", "TARTIŞMA")

    add_body_text(doc,
        "Elde edilen bulgular, FlowGenAI modelinin literatürde sıkça karşılaşılan \"eksik "
        "alternatif akış\" ve \"LLM halüsinasyonları\" problemlerine karşı etkili bir "
        "mühendislik çözümü sunduğunu kanıtlamaktadır. Kural tabanlı NLP sistemlerinin "
        "esneklikten uzak yapısı ile standart LLM'lerin kontrolsüz üretkenliği arasındaki "
        "denge, önerilen melez mimari sayesinde başarıyla kurulmuştur."
    )
    add_body_text(doc,
        "Bölüm II'de özetlenen literatür bulgularıyla karşılaştırıldığında, FlowGenAI'nin "
        "iki temel farklılaşma noktası öne çıkmaktadır: (i) metni yalnızca dönüştürmek "
        "yerine, metnin ötesine geçerek örtük bilgiyi çıkarması; (ii) LLM çıktısını kural "
        "tabanlı doğrulama katmanlarıyla filtrelerek güvenilirliği sistematik biçimde artırması. "
        "%40'lık ek senaryo tespiti, yazılım yaşam döngüsünde tasarım aşamasında fark "
        "edilmeyen hataların ilerleyen fazlarda yaratacağı teknik borcun ciddi oranda "
        "engellendiğini göstermektedir."
    )
    add_body_text(doc,
        "Mevcut sistemin temel kısıtlamaları şu şekilde özetlenebilir: Birincisi, "
        "aşırı derecede çelişkili veya zayıf ifade edilmiş ham gereksinim metinlerinde "
        "kural tabanlı filtrenin güvenlik amacıyla bazı potansiyel alternatif akışları "
        "eleme riski (false-negative üretmesi) bulunmaktadır. İkincisi, sistem "
        "değerlendirmesi yalnızca yazılım mühendisliği alanına ait dokümanlar üzerinde "
        "yapılmış olup farklı sektörlerdeki genellenebilirlik soruşturma gerektirmektedir. "
        "Üçüncüsü, 50 doküman içeren değerlendirme kümesi istatistiksel güç açısından "
        "sınırlı kalabilmektedir."
    )
    add_body_text(doc,
        "Geçerlik tehditler (threats to validity) açısından değerlendirildiğinde, "
        "değerlendirme için seçilen dokümanların sistematik örnekleme yerine erişilebilirlik "
        "kriteri temel alınarak seçilmesi seçim yanlılığı (selection bias) riski "
        "taşımaktadır. Referans değerlendirmelerin üç bağımsız analist tarafından "
        "gerçekleştirilmesi bu riski kısmen azaltmaktadır. Ayrıca, analiz hızı "
        "ölçümlerinde donanım ortamının sabitlenmesi (CPU: Intel Core i7, RAM: 16 GB, "
        "API gecikme ortalaması: 2.1 sn) ölçüm tutarlılığını desteklemektedir."
    )

    # ─── Section VI: SONUÇ ───────────────────────────────────────────────────
    add_section_heading(doc, "VI", "SONUÇ")

    add_body_text(doc,
        "Bu çalışma, yazılım mühendisliği gereksinim analiz süreçlerini optimize etmek "
        "amacıyla tahminleyici ve melez bir yapay zeka modeli olan FlowGenAI'yi sunmaktadır. "
        "Geleneksel yöntemlere kıyasla analiz süresini 10× hızlandıran ve alternatif akış "
        "tespitini %40 oranında genişleten sistem, dört katmanlı doğrulama mekanizması "
        "sayesinde %92 geçerlilik oranıyla güvenilir çıktılar üretmektedir."
    )
    add_body_text(doc,
        "FlowGenAI'nin operasyonel etkisi yalnızca süre azaltımıyla sınırlı değildir; "
        "sistem aynı zamanda analistler arasında tutarlı bir use case kalitesi standardı "
        "oluşturmakta ve teknik borcun oluşmasını analiz aşamasında önlemektedir. "
        "Önerilen melez mimari, mevcut gereksinim mühendisliği araçlarına entegre "
        "edilebilecek modüler bir çerçeve olarak tasarlanmıştır."
    )
    add_body_text(doc,
        "Gelecek çalışmalarda üç temel yön hedeflenmektedir: (i) üretilen use case "
        "yapılarının doğrudan test senaryolarına (test cases) ve otonom kod bloklarına "
        "dönüştürülmesini sağlayan entegrasyonların geliştirilmesi; (ii) değerlendirme "
        "kümesinin farklı sektörleri (finans, sağlık, e-ticaret) kapsayacak biçimde "
        "genişletilerek genellenebilirliğin sınanması; (iii) son-işlemci doğrulama "
        "katmanının aktif öğrenme (active learning) ile uyarlanabilir hale getirilmesi."
    )

    # ─── REFERANSLAR ─────────────────────────────────────────────────────────
    add_section_heading(doc, "", "REFERANSLAR")

    refs = [
        "B. W. Boehm ve V. R. Basili, \"Software defect reduction top 10 list,\" *Computer*, c. 34, s. 1, ss. 135–137, Oca. 2001.",
        "H. Cheng, J. H. Husen, Y. Lu, T. Racharak, N. Yoshioka, N. Ubayashi ve H. Washizaki, \"Generative AI for Requirements Engineering: A Systematic Literature Review,\" *Software: Practice and Experience*, 2025.",
        "J. A. Khan, S. Qayyum ve H. S. Dar, \"Large Language Model for Requirements Engineering: A Systematic Literature Review,\" 2024.",
        "D. K. Deeptimahanti ve R. Sanyal, \"Semi-automatic Generation of UML Models from Natural Language Requirements,\" *ISEC '11*, Thiruvananthapuram, India, Şub. 2011.",
        "N. J. Kamarudin, N. F. M. Sani, R. Atan ve N. C. Pa, \"New Model Transformation Using Requirement Traceability from Requirement to UML Behavioral Design,\" *Journal of Theoretical and Applied Information Technology*, c. 49, s. 3, 2013.",
        "T. H. Nguyen, J. Grundy ve M. Almorsy, \"Rule-Based Extraction of Goal-Use Case Models from Text,\" *ESEC/FSE '15*, Bergamo, Italy, Ağu. 2015.",
        "F. Gilson, M. Galster ve F. Georis, \"Generating Use Case Scenarios from User Stories,\" *ICSSP '20*, Seoul, Republic of Korea, Eki. 2020.",
        "S. Tiwari, D. Ameta ve A. Banerjee, \"An Approach to Identify Use Case Scenarios from Textual Requirements Specification,\" *ISEC '19*, Pune, India, Şub. 2019.",
        "A. M. Sami, Z. Rasheed, M. Waseem, Z. Zhang, H. Tomas ve P. Abrahamsson, \"A Tool for Test Case Scenarios Generation Using Large Language Models,\" *arXiv preprint*, Haz. 2024.",
        "A. M. Sami vd., \"A Tool for Test Case Scenarios Generation Using Large Language Models,\" *arXiv preprint*, 2024.",
        "P. Kogler, A. Falkner ve S. Sperl, \"Reliable Generation of Formal Specifications using Large Language Models,\" *SE 2024 - Companion*, 2024.",
        "A. Owen ve J. Mike, \"Natural Language to Automation: Developing NLP-Based QA Tools to Translate User Stories into Executable Tests,\" 2025.",
        "F. Casati, S. Ceri, S. Paraboschi ve G. Pozzi, \"Specification and Implementation of Exceptions in Workflow Management Systems,\" *ACM Transactions on Database Systems*, c. 24, s. 3, Eyl. 1999.",
        "J. Navarro ve R. Ibarra, \"Automatic test case generation using natural language processing: A systematic mapping study,\" *Information and Software Technology*, c. 189, 2026.",
        "C. Wang, F. Pastore, A. Goknil, L. Briand ve Z. Iqbal, \"Automatic Generation of System Test Cases from Use Case Specifications,\" *ISSTA '15*, Baltimore, MD, USA, Tem. 2015.",
        "C. Arora, M. Sabetzadeh, L. Briand ve F. Zimmer, \"Extracting Domain Models from Natural-Language Requirements: Approach and Industrial Evaluation,\" *ACM/IEEE 19th International Conference on Model Driven Engineering Languages and Systems*.",
        "W. Guo, L. Zhang ve X. Lian, \"Automatically detecting the conflicts between software requirements based on finer semantic analysis,\" *Software Engineering Institute, Beihang University*, 2019.",
    ]

    for i, ref in enumerate(refs):
        add_reference_entry(doc, i + 1, ref)

    # Save
    out_path = r"C:\Users\HP\Desktop\FlowGenAI\FlowGenAI_IEEE_Revised.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    create_paper()
